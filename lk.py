import os
from docx import Document
from docx.shared import Pt
doc = Document()

# Задаем стиль документа
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(20)

run = True
lk = ''
id_edo = ''

# Название организации
ip_or_ooo = str(input('ИП/ООО: '))

# Ввод переменных
while run:
	choose_lk_run = True
	choose_lk = str(input('Для какого сайта данные: platforma(1); evotor(2); sigma(3); taxcom(4): '))

	while True:
		pass
		# Выбор сайта
		if choose_lk == '1':
			lk = 'Сайт: platformaofd.ru'
			break
		elif choose_lk == '2':
			lk = 'Сайт: market.evotor.ru'
			break
		elif choose_lk == '3':
			lk = 'Сайт: cloud.sigma.ru'
			break
		elif choose_lk == '4':
			lk = 'ЭДО ТАКСКОМ.\nСАЙТ: invoice.taxcom.ru'
			id_edo = str(input('ID ЭДО: '))
			break
		else:
			print('Введите значение!!!')

	# Ввод логина и пароля
	login = str(input('Введите Логин: '))
	pwd = str(input('Введите пароль: '))

	# Запись аобзацев в документ
	doc.add_paragraph(lk)
	doc.add_paragraph('Логин: ' + login)
	doc.add_paragraph('Пароль: ' + pwd)
	if id_edo != '':
		doc.add_paragraph('ID: '+ id_edo)
		id_edo = ''
	doc.add_paragraph('\n')

	while True:
		choose_else = str(input('Добавить еще сайт? Y/N: '))

		if choose_else.lower() == 'y':
			run = True
			break
		elif choose_else.lower() == 'n':
			run = False
			break
		else:
			print('Введите Y или N!')

# Сохранение документа
doc.save('g:/__ОФД/'+ip_or_ooo+'.docx')

# Отправка на печать
while True:
	choose_print = str(input('Отправить на печать? Y/N '))

	if choose_print.lower() == 'y':
		os.startfile('g:/__ОФД/'+ip_or_ooo+'.docx', 'print')
	elif choose_print.lower() == 'n':
		break
	else:
		print('Введите Y или N!')
